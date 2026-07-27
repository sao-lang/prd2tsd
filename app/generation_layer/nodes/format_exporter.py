"""FormatExporterNode — 多格式导出（Markdown / HTML / PDF / DOCX）。"""

from __future__ import annotations

import logging
from io import BytesIO

from app.generation_layer.models import GenerationState

logger = logging.getLogger("prd2tsd.format_exporter")


def _md_to_html(md_text: str) -> str:
    """Markdown → HTML 转换。"""
    try:
        import markdown as md_lib

        return md_lib.markdown(
            md_text,
            extensions=["fenced_code", "tables", "codehilite", "toc"],
        )
    except ImportError:
        logger.warning("markdown 库未安装，HTML 导出跳过")
        return f"<pre>{md_text}</pre>"


def _html_to_pdf(html_text: str) -> bytes | None:
    """HTML → PDF 转换。

    优先使用 weasyprint（Linux 生产环境），降级到 fpdf2（纯 Python）。
    """
    # 尝试 weasyprint（质量更高）
    try:
        import weasyprint

        return weasyprint.HTML(string=html_text).write_pdf()  # type: ignore[no-any-return]
    except (ImportError, OSError) as exc:
        logger.debug("weasyprint 不可用，降级到 fpdf2: %s", exc)
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Courier", size=8)

        # 简单文本写入（逐行处理，避免 HTML 标签）
        import html as html_mod
        import re

        text = html_mod.unescape(re.sub(r"<[^>]+>", "", html_text))
        for line in text.split("\n"):
            # 用 latin-1 可编码的字符替代特殊字符
            safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
            pdf.cell(0, 4, safe_line, new_x="LMARGIN", new_y="NEXT")

        buf = BytesIO()
        pdf.output(buf)
        return buf.getvalue()
    except ImportError:
        logger.warning("fpdf2 库未安装，PDF 导出跳过")
        return None


def _md_to_docx(md_text: str) -> bytes | None:
    """Markdown → DOCX 转换。"""
    try:
        import html as html_mod
        import re

        from docx import Document
        from docx.shared import Inches, Pt

        doc = Document()

        # 按行解析，识别标题和段落
        for line in md_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            # 识别 Markdown 标题
            header_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                doc.add_heading(html_mod.unescape(text), level=level)
                continue

            # 普通段落
            p = doc.add_paragraph()
            run = p.add_run(html_mod.unescape(stripped))
            run.font.size = Pt(11)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except ImportError:
        logger.warning("python-docx 库未安装，DOCX 导出跳过")
        return None


class FormatExporterNode:
    """多格式导出节点：Markdown / HTML / PDF / DOCX。"""

    def run(self, state: GenerationState) -> GenerationState:
        """执行格式导出。

        将技术方案文档导出为 Markdown, HTML, PDF, DOCX 格式。
        PDF 优先使用 weasyprint（Linux 生产环境），降级到 fpdf2（纯 Python）。
        导出内容以 base64 编码字符串存入 export_formats。

        Args:
            state: 当前状态。

        Returns:
            更新后的状态，含 export_formats。
        """
        import base64

        result = state.get("generation_result")
        if not result or not result.content:
            return {**state, "export_formats": {}}

        md_content = result.content
        html_content = _md_to_html(md_content)
        pdf_bytes = _html_to_pdf(html_content)
        docx_bytes = _md_to_docx(md_content)

        export_formats: dict[str, str] = {
            "markdown": md_content,
            "html": html_content,
        }

        if pdf_bytes:
            export_formats["pdf"] = base64.b64encode(pdf_bytes).decode("ascii")
        if docx_bytes:
            export_formats["docx"] = base64.b64encode(docx_bytes).decode("ascii")

        return {
            **state,
            "export_formats": export_formats,
        }
