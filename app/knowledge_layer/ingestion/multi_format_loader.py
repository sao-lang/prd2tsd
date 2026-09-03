"""多格式文档文本提取 — bytes → 可索引文本（Block E B3）。

支持 pdf / csv / docx / md / txt / png / jpg 等常用格式。
本模块负责确定性文本与栅格图片提取；图片 OCR 由异步 image_ocr 模块经 Gateway 完成。

依赖：
- pdf 解析：pypdf
- docx 解析：python-docx
- csv/tsv：Python 标准库 csv
"""

from __future__ import annotations

import csv
import io
import warnings
from dataclasses import dataclass
from pathlib import PurePosixPath
from zipfile import BadZipFile, ZipFile

from app.core.config import settings
from app.core.logger import get_logger

logger = get_logger("prd2tsd.multi_format_loader")

# 可入图的文件扩展名
SUPPORTED_EXTENSIONS = {".md", ".txt", ".csv", ".tsv", ".docx", ".pdf", ".png", ".jpg", ".jpeg"}

# 可直接送入 Vision Provider 的图片扩展名
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}


@dataclass(frozen=True)
class ExtractedImage:
    """从文档中提取、等待 OCR 的图片。"""

    content: bytes
    media_type: str
    source_label: str


def get_ext(filename: str) -> str:
    """提取小写扩展名。

    Args:
        filename: 原始文件名。

    Returns:
        小写扩展名（含点），无扩展名时返回 ".bin"。
    """
    idx = filename.rfind(".")
    return filename[idx:].lower() if idx >= 0 else ".bin"


def is_indexable(filename: str) -> bool:
    """文件是否为可入图格式。

    Args:
        filename: 原始文件名。

    Returns:
        是否可入图。
    """
    return get_ext(filename) in SUPPORTED_EXTENSIONS


def extract_text(content: bytes, filename: str) -> str:
    """从文件字节内容提取可索引文本。

    Args:
        content: 文件字节数据。
        filename: 原始文件名（用于判断格式）。

    Returns:
        提取的文本（图片返回元数据占位文本）。

    Raises:
        ValueError: 不支持的格式或提取失败。
    """
    ext = get_ext(filename)

    if ext in IMAGE_EXTENSIONS:
        return _image_placeholder(filename, content)
    if ext in (".md", ".txt"):
        return content.decode("utf-8", errors="replace")
    if ext in (".csv", ".tsv"):
        return _extract_csv(content, ext)
    if ext == ".docx":
        return _extract_docx(content)
    if ext == ".pdf":
        return _extract_pdf(content)

    raise ValueError(f"不支持的格式: {ext}")


def extract_images(content: bytes, filename: str) -> list[ExtractedImage]:
    """提取独立图片或 PDF/DOCX 内嵌栅格图片。

    返回前统一校验图片可解码，并执行单图、总量和数量上限。发现无法解析的
    图片会显式失败，避免文档被标记为 indexed 却静默丢失图片语义。
    """
    if not settings.KNOWLEDGE_OCR_ENABLED:
        return []

    ext = get_ext(filename)
    if ext in IMAGE_EXTENSIONS:
        candidates = [(content, filename)]
    elif ext == ".docx":
        candidates = _extract_docx_images(content)
    elif ext == ".pdf":
        candidates = _extract_pdf_images(content)
    else:
        return []

    if len(candidates) > settings.KNOWLEDGE_OCR_MAX_IMAGES:
        raise ValueError("文档图片数量超过 OCR 限制")
    raw_total_bytes = sum(len(raw) for raw, _ in candidates)
    if raw_total_bytes > settings.KNOWLEDGE_OCR_MAX_TOTAL_IMAGE_BYTES:
        raise ValueError("文档图片总大小超过 OCR 限制")

    images: list[ExtractedImage] = []
    total_bytes = 0
    seen_hashes: set[str] = set()
    for raw, source_label in candidates:
        if len(raw) > settings.KNOWLEDGE_OCR_MAX_IMAGE_BYTES:
            raise ValueError(f"图片 {source_label} 超过 OCR 单图大小限制")
        image = _normalise_image(raw, source_label)
        image_hash = _content_hash(image.content)
        if image_hash in seen_hashes:
            continue
        seen_hashes.add(image_hash)
        if len(image.content) > settings.KNOWLEDGE_OCR_MAX_IMAGE_BYTES:
            raise ValueError(f"图片 {source_label} 超过 OCR 单图大小限制")
        total_bytes += len(image.content)
        if total_bytes > settings.KNOWLEDGE_OCR_MAX_TOTAL_IMAGE_BYTES:
            raise ValueError("文档图片总大小超过 OCR 限制")
        images.append(image)
        if len(images) > settings.KNOWLEDGE_OCR_MAX_IMAGES:
            raise ValueError("文档图片数量超过 OCR 限制")
    return images


def _content_hash(content: bytes) -> str:
    """返回图片内容指纹。"""
    import hashlib

    return hashlib.sha256(content).hexdigest()


def _normalise_image(content: bytes, source_label: str) -> ExtractedImage:
    """验证图片并转换为 Vision API 稳定支持的 PNG/JPEG。"""
    from PIL import Image, UnidentifiedImageError

    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(io.BytesIO(content)) as image:
                image.load()
                image_format = (image.format or "").upper()
                if image_format in {"JPEG", "PNG"}:
                    media_type = "image/jpeg" if image_format == "JPEG" else "image/png"
                    return ExtractedImage(content=content, media_type=media_type, source_label=source_label)

                converted = io.BytesIO()
                target = image.convert("RGBA" if "A" in image.getbands() else "RGB")
                target.save(converted, format="PNG")
                return ExtractedImage(
                    content=converted.getvalue(),
                    media_type="image/png",
                    source_label=source_label,
                )
    except (
        Image.DecompressionBombError,
        Image.DecompressionBombWarning,
        UnidentifiedImageError,
        OSError,
        ValueError,
    ) as exc:
        raise ValueError(f"无法解析待 OCR 图片 {source_label}: {exc}") from exc


def _extract_docx_images(content: bytes) -> list[tuple[bytes, str]]:
    """从 DOCX 的 word/media 目录提取图片。"""
    try:
        with ZipFile(io.BytesIO(content)) as archive:
            members = [
                member
                for member in archive.infolist()
                if not member.is_dir() and member.filename.lower().startswith("word/media/")
            ]
            if len(members) > settings.KNOWLEDGE_OCR_MAX_IMAGES:
                raise ValueError("DOCX 图片数量超过 OCR 限制")
            if any(member.file_size > settings.KNOWLEDGE_OCR_MAX_IMAGE_BYTES for member in members):
                raise ValueError("DOCX 包含超过 OCR 单图大小限制的图片")
            if sum(member.file_size for member in members) > settings.KNOWLEDGE_OCR_MAX_TOTAL_IMAGE_BYTES:
                raise ValueError("DOCX 图片总大小超过 OCR 限制")
            return [(archive.read(member), f"DOCX 图片 {PurePosixPath(member.filename).name}") for member in members]
    except (BadZipFile, OSError, RuntimeError) as exc:
        raise ValueError(f"DOCX 图片提取失败: {exc}") from exc


def _extract_pdf_images(content: bytes) -> list[tuple[bytes, str]]:
    """使用 pypdf 提取每页 XObject 中的栅格图片。"""
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(content))
        images: list[tuple[bytes, str]] = []
        total_bytes = 0
        for page_number, page in enumerate(reader.pages, start=1):
            for image_number, image in enumerate(page.images, start=1):
                if len(images) >= settings.KNOWLEDGE_OCR_MAX_IMAGES:
                    raise ValueError("PDF 图片数量超过 OCR 限制")
                name = PurePosixPath(image.name or f"image-{image_number}").name
                image_data = image.data
                if len(image_data) > settings.KNOWLEDGE_OCR_MAX_IMAGE_BYTES:
                    raise ValueError(f"PDF 图片 {name} 超过 OCR 单图大小限制")
                total_bytes += len(image_data)
                if total_bytes > settings.KNOWLEDGE_OCR_MAX_TOTAL_IMAGE_BYTES:
                    raise ValueError("PDF 图片总大小超过 OCR 限制")
                images.append((image_data, f"PDF 第 {page_number} 页图片 {name}"))
        return images
    except (OSError, RuntimeError, ValueError) as exc:
        raise ValueError(f"PDF 图片提取失败: {exc}") from exc


def _extract_csv(content: bytes, ext: str) -> str:
    """CSV/TSV → 每行转自然语言句子（仅行级文本转换）。

    Args:
        content: 文件字节数据。
        ext: 扩展名（.csv / .tsv）。

    Returns:
        每行一条自然语言记录的文本。
    """
    delimiter = "\t" if ext == ".tsv" else ","
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    lines: list[str] = []
    for row in reader:
        cells = [cell.strip() for cell in row if cell and cell.strip()]
        if cells:
            lines.append("记录: " + "，".join(cells) + "。")
    return "\n".join(lines)


def _extract_docx(content: bytes) -> str:
    """docx → 段落 + 表格文本。

    Args:
        content: 文件字节数据。

    Returns:
        段落与表格行文本。
    """
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append("表格行: " + "，".join(cells))
    return "\n".join(parts)


def _extract_pdf(content: bytes) -> str:
    """pdf → 逐页提取文本。

    Args:
        content: 文件字节数据。

    Returns:
        逐页文本（每页标注页码）。
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[第 {i} 页]\n{page_text.strip()}")
    return "\n\n".join(pages)


def _image_placeholder(filename: str, content: bytes) -> str:
    """图片 → 元数据占位 chunk（可被文件名检索）。

    Args:
        filename: 原始文件名。
        content: 文件字节数据。

    Returns:
        元数据占位文本。
    """
    ext = get_ext(filename)
    size_kb = len(content) / 1024
    return f"[图片: {filename}, 类型 {ext.lstrip('.')}, 大小 {size_kb:.1f}KB]"
