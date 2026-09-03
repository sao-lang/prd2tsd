"""多格式文档文本提取（Block E B3）单元测试。"""

from __future__ import annotations

from io import BytesIO

import pytest

from app.knowledge_layer.ingestion.multi_format_loader import (
    SUPPORTED_EXTENSIONS,
    extract_images,
    extract_text,
    get_ext,
    is_indexable,
)


class TestMultiFormatLoader:
    """multi_format_loader 各格式提取正确性。"""

    def test_get_ext(self) -> None:
        """验证扩展名提取。"""
        assert get_ext("a.md") == ".md"
        assert get_ext("a.PDF") == ".pdf"
        assert get_ext("noext") == ".bin"

    def test_is_indexable(self) -> None:
        """验证可入图格式判断。"""
        assert is_indexable("a.pdf") is True
        assert is_indexable("a.csv") is True
        assert is_indexable("a.png") is True
        assert is_indexable("a.exe") is False
        assert {
            ".md",
            ".txt",
            ".csv",
            ".tsv",
            ".docx",
            ".pdf",
            ".png",
            ".jpg",
            ".jpeg",
        } == SUPPORTED_EXTENSIONS

    def test_extract_markdown(self) -> None:
        """验证 md 提取。"""
        text = extract_text("# 标题\n正文内容".encode(), "doc.md")
        assert "# 标题" in text
        assert "正文内容" in text

    def test_extract_txt(self) -> None:
        """验证 txt 提取。"""
        text = extract_text("纯文本内容".encode(), "note.txt")
        assert "纯文本内容" in text

    def test_extract_csv_rows(self) -> None:
        """验证 CSV 每行转自然语言句子。"""
        content = b"name,age,city\nAlice,30,Beijing\nBob,25,Shanghai\n"
        text = extract_text(content, "data.csv")
        assert "记录: name，age，city。" in text
        assert "记录: Alice，30，Beijing。" in text
        assert "记录: Bob，25，Shanghai。" in text

    def test_extract_tsv_rows(self) -> None:
        """验证 TSV 行转换。"""
        content = b"name\tage\nAlice\t30\n"
        text = extract_text(content, "data.tsv")
        assert "记录: name，age。" in text
        assert "记录: Alice，30。" in text

    def test_extract_docx(self) -> None:
        """验证 docx 段落 + 表格提取。"""
        from docx import Document

        doc = Document()
        doc.add_paragraph("第一个段落")
        doc.add_paragraph("第二个段落")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        buf = BytesIO()
        doc.save(buf)

        text = extract_text(buf.getvalue(), "report.docx")
        assert "第一个段落" in text
        assert "第二个段落" in text
        assert "表格行: A，B" in text

    def test_extract_pdf(self) -> None:
        """验证 pdf 逐页提取（用 fpdf2 生成测试 PDF）。"""
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        pdf.cell(0, 10, "Hello PDF Content")
        pdf.output(buf := BytesIO())

        text = extract_text(buf.getvalue(), "report.pdf")
        assert "Hello PDF Content" in text
        assert "[第 1 页]" in text

    def test_extract_image_placeholder(self) -> None:
        """验证图片返回元数据占位。"""
        text = extract_text(b"\x89PNG fake bytes", "photo.png")
        assert "[图片: photo.png" in text
        assert "类型 png" in text

    def test_extract_jpg_placeholder(self) -> None:
        """验证 jpg 元数据占位。"""
        text = extract_text(b"fake-jpeg", "pic.jpg")
        assert "类型 jpg" in text

    def test_extract_standalone_image_for_ocr(self) -> None:
        """独立图片应作为真实 Vision 输入提取，而不只有元数据占位。"""
        from PIL import Image

        buf = BytesIO()
        Image.new("RGB", (32, 16), "white").save(buf, format="PNG")

        images = extract_images(buf.getvalue(), "architecture.png")

        assert len(images) == 1
        assert images[0].media_type == "image/png"
        assert images[0].source_label == "architecture.png"
        assert images[0].content.startswith(b"\x89PNG")

    def test_extract_docx_embedded_image_for_ocr(self) -> None:
        """DOCX 的 word/media 图片应被提取并保留来源。"""
        from docx import Document
        from PIL import Image

        image_buf = BytesIO()
        Image.new("RGB", (32, 16), "white").save(image_buf, format="PNG")
        doc = Document()
        doc.add_picture(BytesIO(image_buf.getvalue()))
        doc_buf = BytesIO()
        doc.save(doc_buf)

        images = extract_images(doc_buf.getvalue(), "architecture.docx")

        assert len(images) == 1
        assert images[0].source_label == "DOCX 图片 image1.png"

    def test_extract_pdf_embedded_image_for_ocr(self) -> None:
        """PDF 页内图片应被提取并带页码来源。"""
        from fpdf import FPDF
        from PIL import Image

        image_buf = BytesIO()
        Image.new("RGB", (32, 16), "white").save(image_buf, format="PNG")
        pdf = FPDF()
        pdf.add_page()
        pdf.image(BytesIO(image_buf.getvalue()), x=10, y=10, w=20)
        pdf_buf = BytesIO()
        pdf.output(pdf_buf)

        images = extract_images(pdf_buf.getvalue(), "architecture.pdf")

        assert len(images) == 1
        assert images[0].source_label.startswith("PDF 第 1 页图片")

    def test_pdf_image_limit_is_enforced_during_extraction(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """PDF 应边提取边限量，避免先把全部解码图片堆入内存。"""
        from fpdf import FPDF
        from PIL import Image

        from app.core.config import settings

        first = BytesIO()
        second = BytesIO()
        Image.new("RGB", (32, 16), "white").save(first, format="PNG")
        Image.new("RGB", (32, 16), "black").save(second, format="PNG")
        pdf = FPDF()
        pdf.add_page()
        pdf.image(BytesIO(first.getvalue()), x=10, y=10, w=20)
        pdf.image(BytesIO(second.getvalue()), x=40, y=10, w=20)
        pdf_buf = BytesIO()
        pdf.output(pdf_buf)
        monkeypatch.setattr(settings, "KNOWLEDGE_OCR_MAX_IMAGES", 1)

        with pytest.raises(ValueError, match="PDF 图片数量超过 OCR 限制"):
            extract_images(pdf_buf.getvalue(), "too-many-images.pdf")

    def test_invalid_image_is_rejected_for_ocr(self) -> None:
        """伪造扩展名的图片不能送入 Vision Provider。"""
        with pytest.raises(ValueError, match="无法解析待 OCR 图片"):
            extract_images(b"not-an-image", "fake.png")

    def test_unsupported_extension(self) -> None:
        """验证不支持格式抛错。"""
        with pytest.raises(ValueError, match="不支持的格式"):
            extract_text(b"data", "file.exe")
